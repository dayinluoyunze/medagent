# -*- coding: utf-8 -*-
"""
RAG retriever backed by local knowledge files and a persisted FAISS index.
"""

import csv
import hashlib
import json
import os
import plistlib
import re
from collections import Counter
from pathlib import Path
from typing import Iterable, List
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import (
    APP_CONFIG,
    EMBEDDING_PROVIDER,
    KNOWLEDGE_DIR,
    SUPPORTED_KNOWLEDGE_EXTENSIONS,
    VECTORSTORE_DIR,
    VECTORSTORE_MANIFEST,
    get_api_key_for_provider,
    get_provider_config,
)
from rag.ocr import IMAGE_KNOWLEDGE_EXTENSIONS, image_file_to_text, pdf_file_to_ocr_text

try:
    from bs4 import BeautifulSoup
except ImportError:  # pragma: no cover
    BeautifulSoup = None

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover
    DocxDocument = None

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None

try:
    import jieba
except ImportError:  # pragma: no cover
    jieba = None


GENERIC_QUERY_TERMS = {
    "什么",
    "哪些",
    "有哪些",
    "有没有",
    "是否",
    "吗",
    "怎么",
    "怎么办",
    "可以",
    "能否",
    "能不能",
    "应该",
    "需要",
    "注意",
    "注意事项",
    "常见",
    "通用",
    "相关",
    "问题",
    "情况",
    "处理",
    "患者",
    "用药",
    "药物",
    "药品",
    "服药",
    "服用",
    "吃",
    "使用",
}

GENERIC_EXPANSION_GROUPS: list[set[str]] = [
    {"忘记", "漏服", "补服", "下次", "加倍"},
    {"饭前", "饭后", "餐前", "餐后", "随餐", "空腹", "服药时间", "什么时候", "早晨", "夜间", "晚上"},
    {"副作用", "不良反应", "不适", "反应", "风险"},
    {"孕妇", "妊娠", "哺乳", "备孕"},
    {"儿童", "孩子", "小孩"},
    {"老人", "老年人", "老年"},
    {"适合", "适应症", "用于", "人群", "病人"},
    {"剂量", "一天", "几次", "每日", "一次", "用法用量", "起始剂量"},
    {"饮酒", "酒精", "喝酒"},
]

GENERIC_INTENT_TERMS = set().union(*GENERIC_EXPANSION_GROUPS)


class KnowledgeRetriever:
    """Loads local knowledge files and provides vector or keyword retrieval."""

    def __init__(
        self,
        embedding_provider: str | None = None,
        embedding_api_key: str = "",
        knowledge_dir: str | Path = KNOWLEDGE_DIR,
        knowledge_base: str = "medical",
        index_namespace: str | None = None,
    ):
        self.embedding_provider = embedding_provider or EMBEDDING_PROVIDER
        self.embedding_api_key = (
            embedding_api_key
            or get_api_key_for_provider(self.embedding_provider, for_embedding=True)
        )
        self.knowledge_dir = Path(knowledge_dir)
        self.knowledge_base = knowledge_base
        self.index_namespace = index_namespace or self._safe_index_namespace(
            str(self.knowledge_dir)
        )
        self.vectorstore = None
        self.documents: List[Document] = []
        self.chunks: List[Document] = []
        self.init_error = ""
        self.index_dir = os.path.join(
            VECTORSTORE_DIR,
            self.embedding_provider,
            self.index_namespace,
        )
        self.manifest_path = os.path.join(self.index_dir, VECTORSTORE_MANIFEST)
        self._init_vectorstore()

    def _safe_index_namespace(self, value: str) -> str:
        normalized = re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("._-")
        return normalized or "knowledge"

    def _metadata_with_knowledge_base(self, metadata: dict) -> dict:
        enriched = dict(metadata)
        enriched.setdefault("knowledge_base", getattr(self, "knowledge_base", "medical"))
        enriched.setdefault(
            "knowledge_dir",
            str(getattr(self, "knowledge_dir", Path(KNOWLEDGE_DIR))),
        )
        return enriched

    def _allow_remote_url(self, url: str) -> bool:
        if not APP_CONFIG["allow_remote_knowledge_fetch"]:
            return False

        allowlist = APP_CONFIG["remote_knowledge_allowlist"]
        if not allowlist:
            return True

        hostname = (urlparse(url).netloc or "").lower()
        return any(hostname == host.lower() or hostname.endswith(f".{host.lower()}") for host in allowlist)

    def _get_embedding_model(self) -> OpenAIEmbeddings:
        config = get_provider_config(self.embedding_provider, for_embedding=True)
        return OpenAIEmbeddings(
            model=config["model"],
            api_key=self.embedding_api_key,
            base_url=config["api_base"],
            request_timeout=APP_CONFIG["request_timeout"],
            max_retries=APP_CONFIG["max_retries"],
            model_kwargs=config.get("model_kwargs", {}),
        )

    def _knowledge_files(self) -> List[Path]:
        root = Path(getattr(self, "knowledge_dir", KNOWLEDGE_DIR))
        if not root.exists():
            return []

        files: List[Path] = []
        for path in root.rglob("*"):
            if path.is_file() and path.suffix.lower() in SUPPORTED_KNOWLEDGE_EXTENSIONS:
                files.append(path)
        files.sort()
        return files

    def _make_document(self, content: str, filepath: Path, file_type: str) -> Document:
        return Document(
            page_content=content,
            metadata=self._metadata_with_knowledge_base(
                {"source": str(filepath), "file_type": file_type}
            ),
        )

    def _split_frontmatter(self, content: str) -> tuple[dict, str]:
        if not content.startswith("---\n"):
            return {}, content

        end_index = content.find("\n---", 4)
        if end_index < 0:
            return {}, content

        metadata: dict[str, str] = {}
        frontmatter = content[4:end_index]
        for line in frontmatter.splitlines():
            if ":" not in line:
                continue
            key, value = line.split(":", 1)
            metadata[key.strip()] = value.strip().strip("\"'")

        body = content[end_index + 4 :].lstrip()
        return metadata, body

    def _load_text_file(self, filepath: Path) -> List[Document]:
        content = filepath.read_text(encoding="utf-8", errors="ignore")
        frontmatter, body = self._split_frontmatter(content)
        metadata = self._metadata_with_knowledge_base(
            {"source": str(filepath), "file_type": filepath.suffix.lower()}
        )

        source_url = str(frontmatter.get("medagent_source_url", "")).strip()
        if source_url.startswith(("http://", "https://")):
            metadata["source"] = source_url
            metadata["source_file"] = str(filepath)
            metadata["source_title"] = str(frontmatter.get("medagent_source_title", "")).strip()
            metadata["source_type"] = str(frontmatter.get("medagent_source_type", "")).strip()

        return [Document(page_content=body, metadata=metadata)]

    def _load_json_file(self, filepath: Path) -> List[Document]:
        content = filepath.read_text(encoding="utf-8")
        if filepath.suffix.lower() == ".jsonl":
            rows = []
            for line in content.splitlines():
                line = line.strip()
                if line:
                    rows.append(json.loads(line))
            normalized = json.dumps(rows, ensure_ascii=False, indent=2)
        else:
            normalized = json.dumps(json.loads(content), ensure_ascii=False, indent=2)
        return [self._make_document(normalized, filepath, filepath.suffix.lower())]

    def _load_csv_file(self, filepath: Path) -> List[Document]:
        lines: List[str] = []
        with filepath.open("r", encoding="utf-8", newline="") as handle:
            reader = csv.reader(handle)
            for row in reader:
                lines.append(", ".join(cell.strip() for cell in row))
        return [self._make_document("\n".join(lines), filepath, filepath.suffix.lower())]

    def _load_docx_file(self, filepath: Path) -> List[Document]:
        if DocxDocument is None:
            raise RuntimeError("python-docx is not installed")

        doc = DocxDocument(str(filepath))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return [self._make_document("\n".join(paragraphs), filepath, filepath.suffix.lower())]

    def _load_pdf_file(self, filepath: Path) -> List[Document]:
        if PdfReader is None:
            raise RuntimeError("pypdf is not installed")

        reader = PdfReader(str(filepath))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append(f"[page {index}]\n{text}")

        content = "\n\n".join(pages).strip()
        if not content:
            content = pdf_file_to_ocr_text(filepath)
            doc = self._make_document(content, filepath, filepath.suffix.lower())
            doc.metadata["ocr"] = True
            return [doc]

        return [self._make_document(content, filepath, filepath.suffix.lower())]

    def _load_image_file(self, filepath: Path) -> List[Document]:
        content = image_file_to_text(filepath)
        doc = self._make_document(content, filepath, filepath.suffix.lower())
        doc.metadata["ocr"] = True
        return [doc]

    def _load_html_file(self, filepath: Path) -> List[Document]:
        content = filepath.read_text(encoding="utf-8", errors="ignore")
        return [self._make_document(self._html_to_text(content), filepath, filepath.suffix.lower())]

    def _extract_urls_from_file(self, filepath: Path) -> List[str]:
        suffix = filepath.suffix.lower()
        content = filepath.read_text(encoding="utf-8", errors="ignore")

        if suffix == ".webloc":
            try:
                data = plistlib.loads(filepath.read_bytes())
                url = data.get("URL")
                return [url] if isinstance(url, str) and url else []
            except Exception:
                pass

        if suffix == ".url":
            urls = re.findall(r"^URL=(.+)$", content, flags=re.MULTILINE)
            if urls:
                return [url.strip() for url in urls if url.strip()]

        urls = []
        for line in content.splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if re.match(r"^https?://", line, flags=re.IGNORECASE):
                urls.append(line)
        return urls

    def _html_to_text(self, html: str) -> str:
        if BeautifulSoup is not None:
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)

        html = re.sub(r"(?is)<(script|style).*?>.*?</\\1>", " ", html)
        html = re.sub(r"(?s)<[^>]+>", " ", html)
        html = re.sub(r"\s+", " ", html)
        return html.strip()

    def _fetch_url_content(self, url: str) -> str:
        if not self._allow_remote_url(url):
            raise RuntimeError(f"remote knowledge host is not allowed: {url}")

        request = Request(
            url,
            headers={"User-Agent": "medagent-rag-loader/1.0"},
        )
        with urlopen(request, timeout=APP_CONFIG["url_fetch_timeout"]) as response:
            charset = response.headers.get_content_charset() or "utf-8"
            html = response.read().decode(charset, errors="ignore")
        return self._html_to_text(html)

    def _load_url_file(self, filepath: Path) -> List[Document]:
        documents: List[Document] = []
        for url in self._extract_urls_from_file(filepath):
            if not self._allow_remote_url(url):
                continue
            try:
                text = self._fetch_url_content(url)
                if not text:
                    continue
                hostname = urlparse(url).netloc or url
                documents.append(
                    Document(
                        page_content=text,
                        metadata=self._metadata_with_knowledge_base(
                            {
                                "source": url,
                                "source_file": str(filepath),
                                "file_type": filepath.suffix.lower(),
                                "host": hostname,
                            }
                        ),
                    )
                )
            except Exception as exc:
                print(f"Failed to fetch URL {url}: {exc}")
        return documents

    def _load_single_file(self, filepath: Path) -> List[Document]:
        suffix = filepath.suffix.lower()
        if suffix in {".json", ".jsonl"}:
            return self._load_json_file(filepath)
        if suffix == ".csv":
            return self._load_csv_file(filepath)
        if suffix == ".docx":
            return self._load_docx_file(filepath)
        if suffix == ".pdf":
            return self._load_pdf_file(filepath)
        if suffix in IMAGE_KNOWLEDGE_EXTENSIONS:
            return self._load_image_file(filepath)
        if suffix in {".html", ".htm"}:
            return self._load_html_file(filepath)
        if suffix in {".url", ".webloc", ".urls"}:
            return self._load_url_file(filepath)
        return self._load_text_file(filepath)

    def _load_documents(self) -> List[Document]:
        documents: List[Document] = []
        for filepath in self._knowledge_files():
            try:
                documents.extend(self._load_single_file(filepath))
            except Exception as exc:
                print(f"Failed to load knowledge file {filepath.name}: {exc}")
        return documents

    def _split_documents(self, documents: List[Document]) -> List[Document]:
        section_documents: List[Document] = []
        for document in documents:
            section_documents.extend(self._split_document_by_markdown_sections(document))

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n## ", "\n### ", "\n\n", "\n", "。", "；", "，", " ", ""],
        )
        chunks = text_splitter.split_documents(section_documents)
        return [chunk for chunk in chunks if self._has_enough_body_text(chunk)]

    def _has_enough_body_text(self, document: Document) -> bool:
        body = re.sub(r"(?m)^#{1,6}\s+.*$", "", document.page_content)
        body = re.sub(r"\s+", "", body)
        return len(body) >= 8

    def _split_document_by_markdown_sections(self, document: Document) -> List[Document]:
        """Keep evidence chunks focused while preserving parent Markdown headings."""

        file_type = str(document.metadata.get("file_type", "")).lower()
        if file_type not in {".md", ".markdown", ".txt", ".text"}:
            return [document]

        lines = document.page_content.splitlines()
        if not any(re.match(r"^#{1,6}\s+", line.strip()) for line in lines):
            return [document]

        sections: List[Document] = []
        heading_stack: dict[int, str] = {}
        body_lines: List[str] = []
        aliases_by_heading = self._extract_markdown_heading_aliases(document.page_content)

        def flush_section() -> None:
            content_lines = [heading_stack[level] for level in sorted(heading_stack)]
            content_lines.extend(body_lines)
            content = "\n".join(line for line in content_lines if line.strip()).strip()
            if not content:
                return
            metadata = dict(document.metadata)
            metadata["section_headings"] = [
                re.sub(r"^#{1,6}\s+", "", heading_stack[level]).strip()
                for level in sorted(heading_stack)
            ]
            metadata["section"] = " > ".join(metadata["section_headings"])
            aliases: list[str] = []
            for heading in metadata["section_headings"]:
                aliases.extend(aliases_by_heading.get(heading, []))
            if aliases:
                deduped_aliases = list(dict.fromkeys(aliases))
                metadata["section_aliases"] = deduped_aliases
                insert_at = len(content_lines)
                content_lines.insert(insert_at, "别名：" + "、".join(deduped_aliases))
                content = "\n".join(line for line in content_lines if line.strip()).strip()
            sections.append(Document(page_content=content, metadata=metadata))

        for raw_line in lines:
            line = raw_line.rstrip()
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
            if heading_match:
                if body_lines:
                    flush_section()
                    body_lines = []
                level = len(heading_match.group(1))
                heading_stack = {
                    existing_level: heading
                    for existing_level, heading in heading_stack.items()
                    if existing_level < level
                }
                heading_stack[level] = line.strip()
                continue

            if line.strip() == "---":
                if body_lines:
                    flush_section()
                    body_lines = []
                continue

            body_lines.append(line)

        if body_lines or heading_stack:
            flush_section()

        return sections or [document]

    def _extract_markdown_heading_aliases(self, content: str) -> dict[str, list[str]]:
        aliases_by_heading: dict[str, list[str]] = {}
        current_h2 = ""
        for raw_line in content.splitlines():
            line = raw_line.strip()
            heading_match = re.match(r"^##\s+(.+)$", line)
            if heading_match:
                current_h2 = heading_match.group(1).strip()
                aliases_by_heading.setdefault(current_h2, [])
                continue
            if not current_h2:
                continue

            alias_match = re.match(
                r"^-\s*(?:\*\*)?(名称|别名|通用名|商品名|药品名|产品名|英文名)(?:\*\*)?\s*[:：]\s*(.+)$",
                line,
            )
            if not alias_match:
                continue
            alias = re.sub(r"[*`]", "", alias_match.group(2)).strip()
            if alias and alias not in aliases_by_heading[current_h2]:
                aliases_by_heading[current_h2].append(alias)
        return aliases_by_heading

    def _hash_files(self, files: Iterable[Path]) -> str:
        digest = hashlib.sha256()
        root = Path(getattr(self, "knowledge_dir", KNOWLEDGE_DIR))
        for filepath in files:
            try:
                relative_path = filepath.relative_to(root)
            except ValueError:
                relative_path = filepath
            digest.update(str(relative_path).encode("utf-8"))
            digest.update(filepath.read_bytes())
        return digest.hexdigest()

    def _current_manifest(self) -> dict:
        files = self._knowledge_files()
        return {
            "embedding_provider": self.embedding_provider,
            "embedding_model": get_provider_config(
                self.embedding_provider, for_embedding=True
            )["model"],
            "embedding_model_kwargs": get_provider_config(
                self.embedding_provider, for_embedding=True
            ).get("model_kwargs", {}),
            "knowledge_base": self.knowledge_base,
            "knowledge_dir": str(self.knowledge_dir),
            "index_namespace": self.index_namespace,
            "retrieval_mode": APP_CONFIG["retrieval_mode"],
            "allow_remote_knowledge_fetch": APP_CONFIG["allow_remote_knowledge_fetch"],
            "remote_knowledge_allowlist": APP_CONFIG["remote_knowledge_allowlist"],
            "ocr_enabled": APP_CONFIG["ocr_enabled"],
            "ocr_lang": APP_CONFIG["ocr_lang"],
            "ocr_dpi": APP_CONFIG["ocr_dpi"],
            "ocr_max_pages": APP_CONFIG["ocr_max_pages"],
            "knowledge_signature": self._hash_files(files),
            "file_count": len(files),
            "loader_version": 9,
        }

    def _read_manifest(self) -> dict | None:
        if not os.path.exists(self.manifest_path):
            return None
        try:
            return json.loads(Path(self.manifest_path).read_text(encoding="utf-8"))
        except Exception:
            return None

    def _write_manifest(self, manifest: dict) -> None:
        os.makedirs(self.index_dir, exist_ok=True)
        Path(self.manifest_path).write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def _load_cached_vectorstore(self, embedding: OpenAIEmbeddings) -> FAISS | None:
        index_file = os.path.join(self.index_dir, "index.faiss")
        store_file = os.path.join(self.index_dir, "index.pkl")
        if not (os.path.exists(index_file) and os.path.exists(store_file)):
            return None
        try:
            return FAISS.load_local(
                self.index_dir,
                embedding,
                allow_dangerous_deserialization=True,
            )
        except Exception as exc:
            print(f"Failed to load cached vectorstore: {exc}")
            return None

    def _init_vectorstore(self) -> None:
        self.documents = self._load_documents()
        if not self.documents:
            self.init_error = "未找到可加载的知识库文件。"
            return

        self.chunks = self._split_documents(self.documents)

        if self.embedding_provider == "none":
            self.init_error = ""
            return

        if not self.embedding_api_key:
            self.init_error = (
                f"缺少 embedding provider '{self.embedding_provider}' 的 API Key，"
                "已退回本地关键词检索。"
            )
            return

        try:
            embedding = self._get_embedding_model()
            current_manifest = self._current_manifest()
            cached_manifest = self._read_manifest()

            if cached_manifest == current_manifest:
                self.vectorstore = self._load_cached_vectorstore(embedding)
                if self.vectorstore is not None:
                    return

            self.vectorstore = FAISS.from_documents(self.chunks, embedding)
            os.makedirs(self.index_dir, exist_ok=True)
            self.vectorstore.save_local(self.index_dir)
            self._write_manifest(current_manifest)
        except Exception as exc:
            self.init_error = f"初始化向量检索失败：{exc}。已退回本地关键词检索。"
            print(self.init_error)

    def _tokenize_query(self, query: str) -> List[str]:
        normalized = query.strip().lower()
        if not normalized:
            return []

        tokens: List[str] = []
        tokens.extend(
            token.strip().lower()
            for token in re.findall(r"[\u4e00-\u9fff]{1,6}[A-Za-z0-9][A-Za-z0-9_-]*", normalized)
            if token.strip()
        )
        if jieba is not None:
            tokens.extend(
                token.strip().lower()
                for token in jieba.lcut(normalized)
                if token.strip()
            )
        else:
            for span in re.findall(r"[\u4e00-\u9fff]+", normalized):
                for size in range(2, min(len(span), 6) + 1):
                    tokens.extend(span[index : index + size] for index in range(0, len(span) - size + 1))

        regex_pattern = r"[A-Za-z0-9_-]+"
        if jieba is None:
            regex_pattern = r"[\u4e00-\u9fff]{1,8}|[A-Za-z0-9_-]+"
        regex_tokens = re.findall(regex_pattern, normalized)
        tokens.extend(regex_tokens)

        deduped: List[str] = []
        seen = set()
        for token in tokens:
            if len(token) <= 1 and not re.match(r"[A-Za-z0-9]", token):
                continue
            if token not in seen:
                seen.add(token)
                deduped.append(token)
        return deduped

    def _keyword_score(self, query: str, content: str, tokens: List[str]) -> int:
        score = 0
        if query.lower() in content:
            score += 10

        for token in tokens:
            occurrences = content.count(token)
            if not occurrences:
                continue
            token_weight = 1
            if len(token) >= 4:
                token_weight = 3
            elif len(token) >= 2:
                token_weight = 2
            score += occurrences * token_weight

        if score == 0:
            char_counter = Counter(content)
            score = sum(char_counter.get(token, 0) for token in tokens if len(token) == 1)

        return score

    def _base_signal_terms(self, query: str) -> List[str]:
        tokens = self._tokenize_query(query)
        terms: List[str] = []
        seen = set()

        def add(term: str) -> None:
            normalized = term.strip().lower()
            if not normalized or normalized in seen or normalized in GENERIC_QUERY_TERMS:
                return
            if len(normalized) <= 1 and not re.match(r"[A-Za-z0-9]", normalized):
                return
            seen.add(normalized)
            terms.append(normalized)

        for token in tokens:
            add(token)

        return terms

    def _signal_terms(self, query: str) -> List[str]:
        terms = self._base_signal_terms(query)
        seen = set(terms)

        def add(term: str) -> None:
            normalized = term.strip().lower()
            if not normalized or normalized in seen or normalized in GENERIC_QUERY_TERMS:
                return
            if len(normalized) <= 1 and not re.match(r"[A-Za-z0-9]", normalized):
                return
            seen.add(normalized)
            terms.append(normalized)

        query_lower = query.lower()
        for expansion_terms in GENERIC_EXPANSION_GROUPS:
            if any(term.lower() in query_lower for term in expansion_terms):
                for term in expansion_terms:
                    add(term)

        return terms

    def _query_specific_terms(
        self,
        query: str,
        base_terms: List[str] | None = None,
    ) -> List[str]:
        terms = base_terms if base_terms is not None else self._base_signal_terms(query)
        return [
            term
            for term in terms
            if term not in GENERIC_INTENT_TERMS and term not in GENERIC_QUERY_TERMS
        ]

    def _required_query_terms(
        self,
        query: str,
        base_terms: List[str] | None = None,
    ) -> List[str]:
        query_specific_terms = self._query_specific_terms(query, base_terms)
        required_terms: List[str] = []
        for term in sorted(query_specific_terms, key=len, reverse=True):
            if any(
                term != kept
                and term in kept
                and re.search(r"[A-Za-z0-9]", kept)
                for kept in required_terms
            ):
                continue
            required_terms.append(term)
        return required_terms

    def _doc_search_text(self, doc: Document) -> str:
        metadata_text = " ".join(
            str(value)
            for key, value in doc.metadata.items()
            if key in {"source", "source_file", "source_title", "section", "section_aliases"}
        )
        return f"{doc.page_content}\n{metadata_text}".lower()

    def _section_search_text(self, doc: Document) -> str:
        values = [
            doc.metadata.get("source_title", ""),
            doc.metadata.get("section", ""),
            doc.metadata.get("section_headings", ""),
            doc.metadata.get("section_aliases", ""),
        ]
        return " ".join(str(value) for value in values).lower()

    def _relevance_features(
        self,
        query: str,
        doc: Document,
        signal_terms: List[str] | None = None,
    ) -> dict[str, float | int | bool | list[str]]:
        terms = signal_terms if signal_terms is not None else self._signal_terms(query)
        base_terms = self._base_signal_terms(query)
        query_specific_terms = self._query_specific_terms(query, base_terms)
        required_query_terms = self._required_query_terms(query, base_terms)
        doc_text = self._doc_search_text(doc)
        section_text = self._section_search_text(doc)
        matched_terms = [term for term in terms if term in doc_text]
        matched_base_terms = [term for term in base_terms if term in doc_text]
        matched_query_specific_terms = [
            term for term in query_specific_terms if term in doc_text
        ]
        matched_required_terms = [
            term for term in required_query_terms if term in doc_text
        ]
        heading_matches = [term for term in terms if term in section_text]
        keyword_score = self._keyword_score(query, doc_text, terms)
        exact_query_match = bool(query.strip() and query.strip().lower() in doc_text)
        score = float(keyword_score)
        score += min(len(matched_terms), 8) * 1.0
        score += min(len(matched_base_terms), 6) * 1.5
        score += min(len(matched_query_specific_terms), 4) * 3.0
        score += min(len(matched_required_terms), 4) * 2.0
        score += min(len(heading_matches), 4) * 1.5
        if exact_query_match:
            score += 8.0
        return {
            "score": score,
            "keyword_score": keyword_score,
            "signal_overlap": len(matched_terms),
            "matched_terms": matched_terms[:8],
            "base_signal_overlap": len(matched_base_terms),
            "query_specific_overlap": len(matched_query_specific_terms),
            "required_query_overlap": len(matched_required_terms),
            "heading_overlap": len(heading_matches),
            "exact_query_match": exact_query_match,
        }

    def _passes_relevance_gate(
        self,
        query: str,
        doc: Document,
        signal_terms: List[str] | None = None,
    ) -> bool:
        features = self._relevance_features(query, doc, signal_terms)
        terms = signal_terms if signal_terms is not None else self._signal_terms(query)
        if not terms:
            return True

        required_query_terms = self._required_query_terms(query)
        if required_query_terms and int(features["required_query_overlap"]) == 0:
            return False

        min_overlap = APP_CONFIG["retrieval_min_signal_overlap"]
        if required_query_terms or len(terms) >= 4:
            min_overlap = max(min_overlap, 2)
        min_score = APP_CONFIG["retrieval_min_relevance_score"]
        return (
            int(features["signal_overlap"]) >= min_overlap
            and float(features["score"]) >= min_score
        )

    def _annotate_relevance(
        self,
        query: str,
        doc: Document,
        signal_terms: List[str] | None = None,
    ) -> Document:
        features = self._relevance_features(query, doc, signal_terms)
        metadata = dict(doc.metadata)
        metadata["relevance_score"] = round(float(features["score"]), 4)
        metadata["signal_overlap"] = int(features["signal_overlap"])
        metadata["matched_terms"] = features["matched_terms"]
        metadata["query_specific_overlap"] = int(features["query_specific_overlap"])
        metadata["required_query_overlap"] = int(features["required_query_overlap"])
        metadata["heading_overlap"] = int(features["heading_overlap"])
        return Document(page_content=doc.page_content, metadata=metadata)

    def _build_excerpt(self, content: str, tokens: List[str]) -> str:
        flat_content = re.sub(r"\s+", " ", content).strip()
        if not flat_content:
            return ""

        for token in sorted(tokens, key=len, reverse=True):
            index = flat_content.lower().find(token.lower())
            if index >= 0:
                half_window = max(20, APP_CONFIG["citation_snippet_length"] // 2)
                start = max(0, index - half_window)
                end = min(len(flat_content), index + len(token) + half_window)
                snippet = flat_content[start:end].strip()
                if start > 0:
                    snippet = "..." + snippet
                if end < len(flat_content):
                    snippet += "..."
                return snippet

        return flat_content[: APP_CONFIG["citation_snippet_length"]].strip()

    def _decorate_docs(
        self,
        query: str,
        docs: List[Document],
        *,
        retrieval_mode: str,
    ) -> List[Document]:
        tokens = self._tokenize_query(query)
        signal_terms = self._signal_terms(query)
        decorated: List[Document] = []
        for doc in docs:
            annotated = self._annotate_relevance(query, doc, signal_terms)
            metadata = dict(annotated.metadata)
            metadata["retrieval_mode"] = retrieval_mode
            metadata["excerpt"] = self._build_excerpt(annotated.page_content, tokens)
            decorated.append(Document(page_content=annotated.page_content, metadata=metadata))
        return decorated

    def _merge_docs(
        self,
        primary_docs: List[Document],
        secondary_docs: List[Document],
        *,
        k: int,
    ) -> List[Document]:
        merged: List[Document] = []
        seen = set()
        for doc in primary_docs + secondary_docs:
            key = self._doc_key(doc)
            if key in seen:
                continue
            seen.add(key)
            merged.append(doc)
            if len(merged) >= k:
                break
        return merged

    def _doc_key(self, doc: Document) -> tuple[str, str, str]:
        section = str(doc.metadata.get("section", ""))
        return (
            str(doc.metadata.get("source", "")),
            str(doc.metadata.get("source_file", "")),
            section or doc.page_content[:200],
        )

    def _rank_hybrid_docs(
        self,
        query: str,
        vector_docs: List[Document],
        keyword_docs: List[Document],
        *,
        k: int,
    ) -> List[Document]:
        signal_terms = self._signal_terms(query)
        scores: dict[tuple[str, str, str], float] = {}
        docs: dict[tuple[str, str, str], Document] = {}

        for rank, doc in enumerate(vector_docs, start=1):
            if not self._passes_relevance_gate(query, doc, signal_terms):
                continue
            key = self._doc_key(doc)
            docs.setdefault(key, doc)
            features = self._relevance_features(query, doc, signal_terms)
            scores[key] = scores.get(key, 0.0) + 2.5 / rank + min(float(features["score"]) / 8.0, 6.0)

        for rank, doc in enumerate(keyword_docs, start=1):
            if not self._passes_relevance_gate(query, doc, signal_terms):
                continue
            key = self._doc_key(doc)
            docs.setdefault(key, doc)
            features = self._relevance_features(query, doc, signal_terms)
            scores[key] = scores.get(key, 0.0) + 1.8 / rank + min(float(features["score"]) / 6.0, 8.0)

        ranked_keys = sorted(scores, key=lambda item: scores[item], reverse=True)
        ranked_docs: List[Document] = []
        for key in ranked_keys[:k]:
            annotated_doc = self._annotate_relevance(query, docs[key], signal_terms)
            metadata = dict(annotated_doc.metadata)
            metadata["rerank_score"] = round(scores[key], 4)
            ranked_docs.append(Document(page_content=annotated_doc.page_content, metadata=metadata))
        return ranked_docs

    def _keyword_search(self, query: str, k: int = 4) -> List[Document]:
        if not self.chunks:
            return []

        signal_terms = self._signal_terms(query)
        if not signal_terms:
            return self.chunks[:k]

        scored_chunks = []
        for doc in self.chunks:
            if not self._passes_relevance_gate(query, doc, signal_terms):
                continue
            features = self._relevance_features(query, doc, signal_terms)
            score = float(features["score"])

            if score > 0:
                scored_chunks.append((score, doc))

        scored_chunks.sort(key=lambda item: item[0], reverse=True)
        selected_docs: List[Document] = []
        seen = set()
        for _, doc in scored_chunks:
            key = self._doc_key(doc)
            if key in seen:
                continue
            seen.add(key)
            selected_docs.append(doc)
            if len(selected_docs) >= k:
                break
        return selected_docs

    def similarity_search(self, query: str, k: int = 4) -> List[Document]:
        retrieval_mode = APP_CONFIG["retrieval_mode"]
        candidate_multiplier = max(APP_CONFIG["retrieval_candidate_multiplier"], 1)
        candidate_k = max(k * candidate_multiplier, k)
        keyword_docs = self._keyword_search(query, k=candidate_k)

        if self.vectorstore is not None and retrieval_mode in {"auto", "vector", "hybrid"}:
            try:
                vector_docs = self.vectorstore.similarity_search(query, k=candidate_k)
                if retrieval_mode == "vector":
                    signal_terms = self._signal_terms(query)
                    gated_vector_docs = [
                        doc
                        for doc in vector_docs
                        if self._passes_relevance_gate(query, doc, signal_terms)
                    ]
                    return self._decorate_docs(query, gated_vector_docs[:k], retrieval_mode="vector")

                merged_docs = self._rank_hybrid_docs(query, vector_docs, keyword_docs, k=k)
                mode = "hybrid_rerank" if keyword_docs else "vector"
                return self._decorate_docs(query, merged_docs, retrieval_mode=mode)
            except Exception as exc:
                # Query-time embedding calls can still fail because of rate limits or
                # provider availability, so keep keyword fallback available.
                self.init_error = f"向量检索查询失败：{exc}。已退回本地关键词检索。"
                print(self.init_error)
                return self._decorate_docs(query, keyword_docs[:k], retrieval_mode="keyword")

        return self._decorate_docs(query, keyword_docs[:k], retrieval_mode="keyword")

    def get_context(self, query: str, k: int = 4) -> str:
        docs = self.similarity_search(query, k)
        if not docs:
            return "未在知识库中找到相关信息。"
        return "\n\n".join(doc.page_content for doc in docs)


def create_retriever(
    embedding_provider: str | None = None,
    embedding_api_key: str = "",
    knowledge_dir: str | Path = KNOWLEDGE_DIR,
    knowledge_base: str = "medical",
    index_namespace: str | None = None,
) -> KnowledgeRetriever:
    return KnowledgeRetriever(
        embedding_provider=embedding_provider,
        embedding_api_key=embedding_api_key,
        knowledge_dir=knowledge_dir,
        knowledge_base=knowledge_base,
        index_namespace=index_namespace,
    )
